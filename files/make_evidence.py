"""Generate synthetic evidence files for the verifier's document cross-check feature.
Run once: python make_evidence.py  ->  creates ./evidence/*"""
import os, json
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from PIL import Image, ImageDraw, ImageFont
from docx import Document as Docx

EV = os.path.join(os.path.dirname(__file__), "evidence")
os.makedirs(EV, exist_ok=True)
BBOX = {}   # filename -> [x0,y0,x1,y1] of the value to box on image evidence
styles = getSampleStyleSheet()
H = ParagraphStyle("H", parent=styles["Title"], fontSize=15, spaceAfter=6)
SUB = ParagraphStyle("SUB", parent=styles["Normal"], fontSize=9, textColor=colors.grey)
BODY = styles["Normal"]

# 1 — Farmers register (xlsx, 47 rows) — VERIFIED case
rows = [{"S.No": i + 1, "Farmer Name": f"Farmer {i+1:02d}", "Village": "Bhadwasi",
         "Training Date": "2025-07-15", "Signature": "(signed)"} for i in range(47)]
pd.DataFrame(rows).to_excel(os.path.join(EV, "farmers_trained_register.xlsx"), index=False)

# 2 — Dairy revenue (pdf) — VALUE discrepancy (evidence Rs 4,86,400 vs reported 5,90,000)
doc = SimpleDocTemplate(os.path.join(EV, "dairy_revenue_statement.pdf"), pagesize=A4)
items = [["Month", "Milk (L)", "Rate", "Amount (Rs)"],
         ["Apr 2025", "9,200", "32", "2,94,400"], ["May 2025", "6,000", "32", "1,92,000"]]
t = Table(items, colWidths=[3 * cm, 3 * cm, 2.5 * cm, 4 * cm])
t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                       ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
                       ("TEXTCOLOR", (0, 0), (-1, 0), colors.white)]))
doc.build([Paragraph("Mahila Dairy Cooperative Society", H),
           Paragraph("Quarterly Sales & Revenue Statement — FY 2025-26 (Q1), Bhadwasi", SUB),
           Spacer(1, 10), t, Spacer(1, 12),
           Paragraph("<b>Total Revenue Generated (Q1): Rs 4,86,400</b>", BODY),
           Spacer(1, 6), Paragraph("Certified by Society Secretary, Bhadwasi.", SUB)])

# 3 — Nadi certificate (pdf) — COUNT discrepancy (three (3) vs reported 4)
doc = SimpleDocTemplate(os.path.join(EV, "nadi_completion_certificate.pdf"), pagesize=A4)
doc.build([Paragraph("Gram Panchayat — Work Completion Certificate", H),
           Paragraph("Watershed Development Programme, FY 2025-26", SUB), Spacer(1, 12),
           Paragraph("This is to certify that <b>three (3) new Nadi structures</b> have been "
                     "constructed and verified on physical inspection dated 12 August 2025 at "
                     "Bhadwasi, Khejarli and Sathin.", BODY),
           Spacer(1, 16), Paragraph("Sarpanch · Junior Engineer (Watershed)", SUB)])

# 6 — Youth placement (pdf) — VALUE discrepancy (52 vs reported 60)
doc = SimpleDocTemplate(os.path.join(EV, "youth_placement_record.pdf"), pagesize=A4)
doc.build([Paragraph("Skilling Programme — Placement Record", H),
           Paragraph("TechSkills Pvt Ltd, Jodhpur — FY 2025-26", SUB), Spacer(1, 12),
           Paragraph("Appointment letters on file confirm that <b>52 youth placed</b> in "
                     "employment during the reporting year. Salary slips attached for verification.", BODY),
           Spacer(1, 16), Paragraph("Placement Officer · TechSkills Pvt Ltd", SUB)])

# 10 — RWH certificate (pdf) — LOCATION discrepancy (Pipar vs reported Sathin)
doc = SimpleDocTemplate(os.path.join(EV, "rwh_completion_certificate.pdf"), pagesize=A4)
doc.build([Paragraph("Gram Panchayat — Completion Certificate", H),
           Paragraph("Rainwater Harvesting Works, FY 2025-26", SUB), Spacer(1, 12),
           Paragraph("Certified that <b>5 rainwater harvesting (RWH) structures</b> have been "
                     "completed at <b>Pipar</b> village and found satisfactory on verification.", BODY),
           Spacer(1, 16), Paragraph("Sarpanch, Gram Panchayat", SUB)])

# 4 — Health camp attendance (png) — image case (review); record the bbox of "128"
img = Image.new("RGB", (760, 430), "white")
d = ImageDraw.Draw(img)
try:
    fb = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 22)
    fm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
except Exception:
    fb = fm = ImageFont.load_default()
d.text((30, 24), "Mobile Health Van — Health Camp Attendance", font=fb, fill="black")
d.text((30, 64), "Village: Pipar    Date: 09-09-2025    Doctor: Dr. S. Meher", font=fm, fill="black")
d.line((30, 98, 730, 98), fill="black", width=2)
val_bbox = None
for i, (lbl, val) in enumerate([("Male patients seen", "71"), ("Female patients seen", "57"),
                                ("Total patients treated", "128"), ("Medicines (kits)", "104")]):
    y = 118 + i * 42
    d.text((40, y), lbl, font=fm, fill="black")
    d.text((600, y), val, font=fm, fill="black")
    if val == "128":
        bb = d.textbbox((600, y), val, font=fm)
        val_bbox = [bb[0] - 6, bb[1] - 4, bb[2] + 6, bb[3] + 4]
img.save(os.path.join(EV, "health_camp_attendance.png"))
BBOX["health_camp_attendance.png"] = val_bbox

# ---------------------------------------------------------------------------
# Broader evidence types required by the technical specification
# ---------------------------------------------------------------------------

# A — Water Holding Capacity Increased (kL): measurement/engineering record (pdf)
#     spec's own example KPI; VALUE discrepancy (1,250 kL vs reported 1,500)
doc = SimpleDocTemplate(os.path.join(EV, "whc_measurement_record.pdf"), pagesize=A4)
doc.build([Paragraph("Watershed Engineering — Measurement Record", H),
           Paragraph("Water Harvesting Structure, Khejarli — FY 2025-26", SUB), Spacer(1, 12),
           Paragraph("On physical measurement and survey, the <b>water holding capacity "
                     "increased by 1,250 kL</b> after desilting and bund strengthening. "
                     "Certified by the Junior Engineer and countersigned by the Panchayat "
                     "representative.", BODY),
           Spacer(1, 16), Paragraph("Junior Engineer (Watershed) · Sarpanch, Gram Panchayat", SUB)])

# B — Exposure tours: field report (Word .docx) — MATCH (verified)
dx = Docx()
dx.add_heading("Quarterly Field Report — Skilling Programme", level=1)
dx.add_paragraph("Implementing partner: TechSkills, Jodhpur. FY 2025-26.")
dx.add_paragraph("During the quarter, 3 exposure tours were conducted for enrolled youth to "
                 "industry sites and training institutes. Travel logs and photographs are on file.")
dx.add_paragraph("Prepared by: Programme Coordinator.")
dx.save(os.path.join(EV, "field_report_exposure_tours.docx"))

# C — Households benefited: register transcription in Odia — MATCH (multilingual)
with open(os.path.join(EV, "household_register_odia.txt"), "w", encoding="utf-8") as f:
    f.write("ଗୃହ ସର୍ବେକ୍ଷଣ ପଞ୍ଜିକା\n"
            "ଗ୍ରାମ: ଯାଜପୁର, ଜିଲ୍ଲା\n"
            "ସର୍ବମୋଟ ଉପକୃତ ପରିବାର: ୩୨୦\n"
            "(Jajpur — Total households benefited: 320.)")

# D — Check dams: photograph of a project signboard (infrastructure) — image (review)
board = Image.new("RGB", (760, 460), "#0F766E")
db = ImageDraw.Draw(board)
try:
    bb1 = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
    bb2 = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
    bb3 = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
except Exception:
    bb1 = bb2 = bb3 = ImageFont.load_default()
db.rectangle([20, 20, 740, 440], outline="white", width=4)
db.text((60, 70), "WATERSHED DEVELOPMENT", font=bb2, fill="white")
db.text((60, 120), "Check Dam Construction", font=bb1, fill="white")
db.text((60, 210), "Units completed:", font=bb3, fill="white")
db.text((360, 205), "2", font=bb1, fill="#FFE08A")
bb = db.textbbox((360, 205), "2", font=bb1)
BBOX["check_dam_signboard.png"] = [bb[0] - 8, bb[1] - 6, bb[2] + 8, bb[3] + 6]
db.text((60, 300), "FY 2025-26", font=bb3, fill="white")
db.text((60, 360), "Gram Panchayat, Sathin", font=bb3, fill="white")
board.save(os.path.join(EV, "check_dam_signboard.png"))

json.dump(BBOX, open(os.path.join(EV, "bbox.json"), "w"))
print("evidence written to", EV)
print("bboxes:", BBOX)
print("files:", sorted(os.listdir(EV)))
